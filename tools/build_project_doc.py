from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"C:\Users\Keso\Downloads\Proyecto CTI.docx")
OUT = ROOT / "docs" / "SOC_Use_Cases_Manager_Formato_CTI.docx"

ACCENT = RGBColor(192, 0, 0)
DARK = RGBColor(0, 0, 0)
MUTED = RGBColor(90, 105, 125)
HEADER_FILL = "C00000"
LIGHT_GRAY = "F4F6F8"
TABLE_BORDER = "808080"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=TABLE_BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_widths(table, widths_in):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx >= len(row.cells):
                continue
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths=None, header=True):
    set_table_borders(table)
    if widths:
        set_table_widths(table, widths)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.08
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    run.font.color.rgb = DARK
            if header and r_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    if "Table Grid" in [s.name for s in doc.styles]:
        table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row_data in rows:
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = str(value)
    style_table(table, widths=widths)
    doc.add_paragraph()
    return table


def set_run(run, size=None, bold=None, color=None, italic=None, font="Arial"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run(run, size={1: 14, 2: 12, 3: 11}.get(level, 10.5), bold=True if level == 1 else None, color=DARK)
    return p


def add_para(doc, text="", *, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run(r1, size=10.5, bold=True, color=DARK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run(r2, size=10.5, color=DARK)
    else:
        run = p.add_run(text)
        set_run(run, size=10.5, color=DARK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run(run, size=10.2, color=DARK)


def add_number(doc, text):
    style = "List Number" if "List Number" in [s.name for s in doc.styles] else "List Paragraph"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    prefix = "" if style == "List Number" else "- "
    run = p.add_run(f"{prefix}{text}")
    set_run(run, size=10.2, color=DARK)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run(run, size=9, color=RGBColor(40, 51, 65), font="Courier New")
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    if "Table Grid" in [s.name for s in doc.styles]:
        table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "FBECEC")
    set_cell_margins(cell, top=140, bottom=140, start=160, end=160)
    set_table_borders(table, "C00000")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run(r, size=10.5, bold=True, color=ACCENT)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run(r2, size=10, color=DARK)
    doc.add_paragraph()


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name in ["List Bullet", "List Number", "List Paragraph"]:
        if name in [s.name for s in styles]:
            styles[name].font.name = "Arial"
            styles[name].font.size = Pt(10.2)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("SOC Use Cases Manager | Documentacion tecnica")
    set_run(run, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Documento tecnico-operativo | Uso interno")
    set_run(r, size=8.5, color=MUTED)


def cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("SOC USE CASES MANAGER")
    set_run(r, size=22, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("Documentacion tecnica, funcional y operativa")
    set_run(r, size=15, color=DARK)

    rows = [
        ("Proyecto", "SOC Use Cases Manager"),
        ("Tipo de documento", "Informe tecnico-operativo"),
        ("Version", "1.0"),
        ("Fecha", date.today().strftime("%d/%m/%Y")),
        ("Repositorio sugerido", "https://github.com/mrkeso1/soc_usecases.git"),
        ("Stack", "Django 6, PostgreSQL, Docker Compose"),
    ]
    add_table(doc, ["Parametro", "Detalle"], rows, widths=[1.75, 4.95])

    add_callout(
        doc,
        "Objetivo del documento",
        "Consolidar en un unico entregable la descripcion funcional, arquitectura, requerimientos tecnicos, integraciones, fuentes externas, cron de sincronizacion, logs, instalacion y operacion del sistema.",
    )

    add_heading(doc, "Resumen ejecutivo", 1)
    add_para(
        doc,
        "SOC Use Cases Manager es una aplicacion web orientada a inventariar, administrar y revisar casos de uso SOC, con cobertura ATT&CK y D3FEND, dashboard ejecutivo, exportaciones, ciclo de vida, roles y autenticacion local/LDAP configurable.",
    )
    for item in [
        "Centraliza el inventario productivo de casos de uso y su mapeo ATT&CK.",
        "Calcula cobertura D3FEND inferida desde relaciones oficiales D3FEND->ATT&CK.",
        "Permite operar lifecycle, responsables, evidencias y proximas revisiones.",
        "Incluye dashboard, PDF ejecutivo, Excel de importacion/exportacion y datos demo.",
        "Automatiza catalogos mediante cron completo de ATT&CK, D3FEND, mappings y casos.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()


def body(doc):
    add_heading(doc, "1. Alcance y objetivos", 1)
    add_para(doc, "El sistema cubre la gestion diaria y ejecutiva de casos de uso SOC. El alcance principal es mantener un inventario confiable, medir cobertura de deteccion frente a frameworks defensivos/ofensivos, y facilitar revisiones periodicas con evidencia historica.")
    add_table(
        doc,
        ["Objetivo", "Descripcion"],
        [
            ("Inventario operativo", "Alta, edicion, filtros, bulk update, importacion y exportacion de casos de uso."),
            ("Cobertura ejecutiva", "Dashboard y PDF con KPIs ATT&CK/D3FEND, pendientes y principales tecnicas."),
            ("Lifecycle", "Asignacion de responsables, finalizacion de controles y calculo de proximas revisiones."),
            ("Integracion de catalogos", "Sincronizacion completa de ATT&CK Enterprise, D3FEND, relaciones y casos inferidos."),
            ("Seguridad de acceso", "Roles Admin, Analyst y ReadOnly, con opcion LDAP/LDAPS administrable."),
        ],
        widths=[2.0, 4.7],
    )

    add_heading(doc, "2. Arquitectura general", 1)
    add_para(doc, "La aplicacion corre como servicio Django dentro de Docker y persiste informacion en PostgreSQL. Las fuentes externas se consultan solo durante procesos de carga/sincronizacion o validacion LDAP.")
    add_code_block(
        doc,
        "Usuario -> Navegador -> Django web:8000 -> PostgreSQL db:5432\n"
        "                         |-> Logs: /logs/*.log\n"
        "                         |-> Media: logos PDF\n"
        "                         |-> HTTPS 443: MITRE ATT&CK / D3FEND\n"
        "                         |-> LDAP/LDAPS: 389/636 si se habilita",
    )
    add_table(
        doc,
        ["Componente", "Responsabilidad"],
        [
            ("apps.accounts", "Usuario custom, roles, LDAPSettings, LDAPAuthLog y backends de autenticacion."),
            ("apps.usecases", "Dominio principal: casos, catalogos, dashboard, matrices, lifecycle, PDF, Excel y cron."),
            ("PostgreSQL", "Persistencia de usuarios, casos, catalogos, relaciones, logs admin y configuraciones."),
            ("Docker Compose", "Orquesta web y db para desarrollo, pruebas y despliegue simple."),
            ("Filesystem", "Almacena logs rotativos y media para logos de reportes."),
        ],
        widths=[1.8, 4.9],
    )

    add_heading(doc, "3. Funcionalidades principales", 1)
    add_heading(doc, "3.1 Inventario de casos", 2)
    for item in [
        "Listado filtrable por texto, dispositivo, severidad, owner, estado, habilitado, revision y mapeo.",
        "Alta/edicion con tabs de datos principales, clasificacion, MITRE/D3FEND, lifecycle y notas.",
        "Bulk update para cambios masivos controlados.",
        "Changelog automatico de campos relevantes.",
        "Importacion Excel desde UI y consola, con plantilla y actualizacion por nombre.",
        "Exportacion Excel/CSV respetando filtros.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.2 Dashboard, matrices y PDF", 2)
    add_para(doc, "El dashboard calcula cobertura sobre casos productivos y reutiliza el mismo contexto para la exportacion PDF. La cobertura D3FEND se infiere desde ATT&CK mediante relaciones oficiales D3FEND->ATT&CK.")
    add_table(
        doc,
        ["Vista/Reporte", "Contenido"],
        [
            ("Dashboard", "KPIs de casos productivos, cobertura ATT&CK, cobertura D3FEND y pendientes."),
            ("Matriz ATT&CK", "Cobertura por tactica, tecnica y subtecnica."),
            ("Matriz D3FEND", "Cobertura defensiva inferida y controles parciales/completos."),
            ("PDF ejecutivo", "Reporte con branding configurable, KPIs, tablas y graficos."),
        ],
        widths=[1.7, 5.0],
    )

    add_heading(doc, "3.3 Lifecycle", 2)
    for item in [
        "Ventanas de revision y proxima fecha configurable por LifecycleSettings.",
        "Responsable de control por caso.",
        "Finalizacion con resultado, evidencia historica y recalculo de proxima revision.",
        "Permisos: Admin puede reasignar; Analyst finaliza casos propios o asignados.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.4 Autenticacion y roles", 2)
    add_table(
        doc,
        ["Rol", "Permisos"],
        [
            ("Admin", "Acceso completo, admin Django, configuraciones, catalogos, coverage admin y lifecycle."),
            ("Analyst", "Gestion de casos propios/asignados, inventario, lifecycle asignado e import/export segun permisos."),
            ("ReadOnly", "Acceso de consulta al dashboard sin operacion de inventario ni admin."),
        ],
        widths=[1.4, 5.3],
    )

    add_heading(doc, "4. Requerimientos tecnicos", 1)
    add_table(
        doc,
        ["Componente", "Requisito"],
        [
            ("Runtime", "Python 3.12+"),
            ("Framework", "Django 6.0.5"),
            ("Base de datos", "PostgreSQL"),
            ("Contenedores", "Docker y Docker Compose"),
            ("Servidor app", "Gunicorn recomendado en produccion; runserver para desarrollo local."),
            ("Sistema recomendado", "Linux para produccion; Windows valido para desarrollo local con Docker Desktop."),
        ],
        widths=[2.0, 4.7],
    )
    add_table(
        doc,
        ["Paquete", "Uso"],
        [
            ("Django", "Framework web y ORM."),
            ("psycopg[binary]", "Driver PostgreSQL."),
            ("gunicorn", "Servidor WSGI."),
            ("openpyxl", "Importacion/exportacion Excel."),
            ("requests", "Descarga de catalogos ATT&CK/D3FEND."),
            ("ldap3", "Autenticacion y pruebas LDAP/LDAPS."),
            ("Pillow", "Imagenes/logos de reportes."),
            ("reportlab", "PDF ejecutivo."),
        ],
        widths=[2.0, 4.7],
    )

    add_heading(doc, "5. Variables de entorno", 1)
    add_table(
        doc,
        ["Variable", "Obligatoria", "Descripcion"],
        [
            ("SECRET_KEY", "Si", "Clave secreta Django; debe ser unica en produccion."),
            ("DEBUG", "Si", "0 en produccion."),
            ("ALLOWED_HOSTS", "Si", "Hosts permitidos separados por coma."),
            ("POSTGRES_DB", "Si", "Nombre de la base."),
            ("POSTGRES_USER", "Si", "Usuario DB."),
            ("POSTGRES_PASSWORD", "Si", "Password DB."),
            ("POSTGRES_HOST", "Si", "En Docker suele ser db."),
            ("POSTGRES_PORT", "Si", "Default interno 5432."),
            ("LOG_DIR", "No", "Directorio de logs; en Docker /logs."),
        ],
        widths=[1.55, 1.1, 4.05],
    )

    add_heading(doc, "6. Red, puertos y conexiones", 1)
    add_table(
        doc,
        ["Servicio", "Docker", "Host local", "Uso"],
        [
            ("web", "8000/tcp", "8000/tcp", "Django, UI, Admin, PDF y Excel."),
            ("db", "5432/tcp", "5433/tcp", "PostgreSQL; app usa db:5432 dentro de Docker."),
        ],
        widths=[1.3, 1.3, 1.3, 2.8],
    )
    add_para(doc, "En produccion se recomienda exponer la aplicacion detras de un reverse proxy HTTPS. El puerto externo final depende del proxy, pero el contenedor web escucha en 8000.")

    add_heading(doc, "7. Fuentes externas y cantidad de consultas", 1)
    add_para(doc, "La sincronizacion completa necesita salida HTTPS por 443/tcp. ATT&CK se obtiene desde la fuente oficial STIX 2.1 publicada por MITRE en GitHub; D3FEND se obtiene desde recursos oficiales MITRE D3FEND.")
    add_table(
        doc,
        ["Fase", "Fuente", "Host:puerto", "Consultas", "Timeout"],
        [
            ("ATT&CK Enterprise", "MITRE ATT&CK STIX 2.1 oficial, repo mitre-attack/attack-stix-data", "raw.githubusercontent.com:443", "1", "120s"),
            ("Catalogo D3FEND", "MITRE D3FEND Ontology Releases", "d3fend.mitre.org:443", "1", "120s"),
            ("Mappings D3FEND->ATT&CK", "MITRE D3FEND API/recurso de relaciones inferidas", "d3fend.mitre.org:443", "2", "120s"),
            ("Normalizacion D3FEND", "MITRE D3FEND API/sitio oficial", "d3fend.mitre.org:443", "Hasta 2 por D3FEND sin codigo D3-*", "30s"),
        ],
        widths=[1.45, 2.65, 1.35, 0.95, 0.6],
    )
    add_code_block(doc, "Consultas externas aproximadas por corrida completa:\n4 + hasta 2 * cantidad_de_d3fend_sin_codigo_oficial")
    add_para(doc, "Nota: attack.mitre.org es el sitio humano de ATT&CK; para sincronizacion estructurada se usa el JSON STIX oficial mantenido por MITRE en GitHub.")

    add_heading(doc, "8. Cron de sincronizacion completa", 1)
    add_para(doc, "La agenda vive en Django Admin, modelo MitreAttackSyncSettings. Una configuracion activa define intervalo en horas o dias, ultimo estado, ultimo mensaje, proxima ejecucion y contadores.")
    add_code_block(
        doc,
        "docker compose run --rm web python manage.py sync_security_frameworks_scheduled\n"
        "docker compose run --rm web python manage.py sync_security_frameworks_scheduled --force",
    )
    for item in [
        "Sincroniza ATT&CK Enterprise.",
        "Carga/actualiza D3FEND.",
        "Reconstruye mappings D3FEND->ATT&CK.",
        "Normaliza codigos D3FEND cuando corresponde.",
        "Recalcula D3FEND inferido en casos de uso.",
    ]:
        add_number(doc, item)
    add_callout(doc, "Recomendacion operativa", "Programar el cron externo cada hora y dejar que MitreAttackSyncSettings decida si corresponde ejecutar segun interval_value e interval_unit.")

    add_heading(doc, "9. LDAP/LDAPS", 1)
    add_table(
        doc,
        ["Modo", "Puerto tipico", "Configuracion"],
        [
            ("LDAP sin TLS", "389/tcp", "server_uri=ldap://servidor:389"),
            ("LDAPS", "636/tcp", "server_uri=ldaps://servidor:636"),
        ],
        widths=[1.8, 1.4, 3.5],
    )
    add_para(doc, "El login LDAP busca el usuario con user_search_base/user_search_filter o arma DN con user_dn_template. Luego intenta bind con usuario y password. El boton Probar conexion hace bind con bind_dn y bind_password y registra el resultado.")

    add_heading(doc, "10. Logs y auditoria", 1)
    add_table(
        doc,
        ["Archivo", "Contenido"],
        [
            ("logs/auth.log", "Login, logout, fallos de login y eventos LDAP."),
            ("logs/mitre_sync.log", "Descarga, omisiones, errores y resultados de sincronizacion MITRE/frameworks."),
            ("logs/app.log", "Warnings y errores HTTP Django."),
        ],
        widths=[2.0, 4.7],
    )
    add_para(doc, "Los logs son rotativos: 5 MB por archivo y 5 backups. Docker monta ./logs del host en /logs dentro del contenedor.")

    add_heading(doc, "11. Instalacion local con Docker", 1)
    add_code_block(
        doc,
        "git clone https://github.com/mrkeso1/soc_usecases.git\n"
        "cd soc_usecases\n"
        "cp .env.example .env\n"
        "docker compose up -d --build\n"
        "docker compose run --rm web python manage.py migrate\n"
        "docker compose run --rm web python manage.py seed_groups\n"
        "docker compose run --rm web python manage.py createsuperuser\n"
        "docker compose run --rm web python manage.py seed_demo_data",
    )
    add_para(doc, "Luego abrir http://localhost:8000/.")

    add_heading(doc, "12. Datos demo y validacion", 1)
    add_para(doc, "El comando seed_demo_data crea usuarios por rol, catalogos ATT&CK/D3FEND, casos productivos/test/desarrollo, revisiones lifecycle, overrides y configuraciones demo.")
    add_table(
        doc,
        ["Usuario", "Rol", "Password default"],
        [
            ("demo_admin", "Admin", "Demo12345!"),
            ("demo_analyst", "Analyst", "Demo12345!"),
            ("demo_owner", "Analyst/control owner", "Demo12345!"),
            ("demo_readonly", "ReadOnly", "Demo12345!"),
        ],
        widths=[2.0, 2.2, 2.5],
    )
    add_code_block(
        doc,
        "docker compose run --rm web python manage.py test --noinput\n"
        "docker compose run --rm web python manage.py makemigrations --check --dry-run",
    )

    add_heading(doc, "13. Checklist de despliegue", 1)
    for item in [
        "Configurar SECRET_KEY, DEBUG=0, ALLOWED_HOSTS, PostgreSQL y LOG_DIR.",
        "Ejecutar migraciones y seed_groups.",
        "Crear o validar superusuario.",
        "Configurar MitreAttackSyncSettings activo.",
        "Configurar cron externo para sync_security_frameworks_scheduled.",
        "Configurar LDAP solo si aplica y probar conexion.",
        "Montar volumen persistente para MEDIA_ROOT si se usan logos PDF.",
        "Validar Dashboard, PDF, Inventario, Lifecycle, LDAP Admin y logs.",
        "Ejecutar tests y makemigrations --check antes de desplegar.",
    ]:
        add_number(doc, item)

    add_heading(doc, "14. Fuentes y referencias", 1)
    add_table(
        doc,
        ["Fuente", "Uso"],
        [
            ("https://github.com/mitre-attack/attack-stix-data", "Dataset oficial ATT&CK STIX 2.1."),
            ("https://raw.githubusercontent.com/mitre-attack/attack-stix-data/master/enterprise-attack/enterprise-attack.json", "JSON Enterprise ATT&CK consumido por el sync."),
            ("https://d3fend.mitre.org/resources/ontology/", "Recursos oficiales de ontologia D3FEND."),
            ("https://d3fend.mitre.org/api/ontology/inference/d3fend-full-mappings.csv", "Relaciones inferidas D3FEND->ATT&CK."),
            ("https://d3fend.mitre.org/api-docs/", "Documentacion API D3FEND."),
            ("https://docs.djangoproject.com/", "Framework Django."),
            ("https://www.postgresql.org/docs/", "Base de datos PostgreSQL."),
        ],
        widths=[3.5, 3.2],
    )

    add_heading(doc, "Anexo A. Comandos operativos", 1)
    add_code_block(
        doc,
        "docker compose run --rm web python manage.py migrate\n"
        "docker compose run --rm web python manage.py seed_groups\n"
        "docker compose run --rm web python manage.py seed_demo_data --reset\n"
        "docker compose run --rm web python manage.py import_usecases archivo.xlsx --update\n"
        "docker compose run --rm web python manage.py load_mitre_attack\n"
        "docker compose run --rm web python manage.py load_d3fend\n"
        "docker compose run --rm web python manage.py sync_security_frameworks_scheduled --force",
    )


def main():
    doc = Document(TEMPLATE) if TEMPLATE.exists() else Document()
    if TEMPLATE.exists():
        body_el = doc._body._element
        for child in list(body_el):
            if child.tag != qn("w:sectPr"):
                body_el.remove(child)
    configure_doc(doc)
    cover(doc)
    body(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
